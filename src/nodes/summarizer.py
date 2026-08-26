import os
import json
from datetime import datetime

from langchain_core.runnables import RunnableConfig
from langchain_core.messages import AIMessage
from langchain_core.prompts import ChatPromptTemplate

from ..state import State
from ..llm import get_llm, invoke_llm


"""
Summarizer 节点（整合自 Code/Summarizer.py）：
- 职责：汇总 Verifier 验证通过的结果，生成 Word 报告（docx），并调用 LLM 生成简明评价；将评价以消息形式追加到状态，并输出报告文件路径。
- 兼容：LangGraph 节点签名 summarizer(state: State, config: RunnableConfig, **kwargs)，与 graph.py 集成。
"""


def _read_prompt(rel_path: str) -> str:
    base_dir = os.path.dirname(__file__)
    path = os.path.join(base_dir, rel_path)
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception:
        return (
            "你是一位专业的报告评价专家。阅读报告全文，按结构/完整性/专业性/总体评价四个维度，用约200字输出自然语言评价。"
        )


def find_title(state: State) -> str:
    for msg in state.get("messages", []) or []:
        try:
            content = json.loads(msg.content)
        except Exception:
            continue
        if content.get("from") == "Intake" and content.get("to") == "Planner":
            return content.get("title") or "自动生成报告"
    return "自动生成报告"


def _content_reorganizer(state: State):
    sections = []
    tasks = state.get("tasks", []) or []
    for res in state.get("results", []) or []:
        # 任务名作为章节标题（若找不到则使用 task_id）
        title = None
        for t in tasks:
            if t.get("task_id") == res.get("task_id"):
                title = t.get("task_name")
                break
        title = title or res.get("task_id") or "章节"

        # 将 outputs 中的图片路径转换为 figures
        figures = []
        for p in (res.get("outputs", []) or []):
            if str(p).lower().endswith((".png", ".jpg", ".jpeg", ".svg")):
                figures.append({"path": p, "description": f"图像：{os.path.basename(p)}"})

        section = {
            "title": title,
            "text": res.get("text_output", ""),
            "tables": res.get("tables", []),
            "figures": figures or res.get("figures", []),
            "citations": res.get("citations", []),
            "notes": res.get("notes", ""),
        }
        sections.append(section)
    return sections


def _docx_report_generator(state: State, sections):
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.shared import RGBColor
    except Exception:
        return None

    def _set_title_style(paragraph, font_size=Pt(16), bold=True):
        for run in paragraph.runs:
            run.font.name = u"宋体"
            run.font.element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.size = font_size
            run.font.bold = bold
        return paragraph

    def _set_normal_style(paragraph):
        for run in paragraph.runs:
            run.font.name = u"宋体"
            run.font.element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.size = Pt(12)
            run.font.bold = False
        return paragraph

    def _add_table_to_doc(doc, table_data, description):
        table_title = doc.add_heading(level=3)
        table_title.add_run(f"表格：{description}")
        _set_title_style(table_title, font_size=Pt(14), bold=True)
        table_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        rows = len(table_data)
        cols = len(table_data[0]) if rows > 0 else 0
        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"

        if rows > 0:
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(table_data[0]):
                hdr_cells[i].text = str(header)
                for paragraph in hdr_cells[i].paragraphs:
                    _set_title_style(paragraph, font_size=Pt(12), bold=True)

        for row_idx in range(1, rows):
            row_cells = table.rows[row_idx].cells
            for col_idx, cell_value in enumerate(table_data[row_idx]):
                row_cells[col_idx].text = str(cell_value)
                for paragraph in row_cells[col_idx].paragraphs:
                    _set_normal_style(paragraph)

        doc.add_paragraph()

    def _add_figure_to_doc(doc, figure_path, description):
        fig_title = doc.add_heading(level=3)
        fig_title.add_run(f"图：{description}")
        _set_title_style(fig_title, font_size=Pt(14), bold=True)
        fig_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if os.path.exists(figure_path):
            doc.add_picture(figure_path, width=Inches(5))
            doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            warning = doc.add_paragraph(f"图片文件不存在：{figure_path}")
            _set_normal_style(warning)
            warning.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph()

    doc = Document()

    title_text = find_title(state)
    title = doc.add_heading(level=1)
    title.add_run(title_text)
    _set_title_style(title, font_size=Pt(20), bold=True)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()

    for sec in sections:
        sec_title = doc.add_heading(level=2)
        sec_title.add_run(sec.get("title", "章节"))
        _set_title_style(sec_title, font_size=Pt(16), bold=True)

        content_para = doc.add_paragraph(sec.get("text", ""))
        _set_normal_style(content_para)
        content_para.space_after = Pt(12)

        for table in sec.get("tables", []) or []:
            data = table.get("data") if isinstance(table, dict) else table
            desc = (table.get("description") if isinstance(table, dict) else "表格")
            if data:
                _add_table_to_doc(doc, data, desc)

        for fig in sec.get("figures", []) or []:
            p = fig.get("path") if isinstance(fig, dict) else fig
            desc = fig.get("description") if isinstance(fig, dict) else os.path.basename(str(p))
            if p:
                _add_figure_to_doc(doc, p, desc)

        if sec.get("notes"):
            note_para = doc.add_paragraph(f"备注：{sec.get('notes')}")
            _set_normal_style(note_para)
            note_para.space_after = Pt(12)

        doc.add_page_break()

    os.makedirs("cache/report", exist_ok=True)
    output_path = os.path.join("cache/report", "report.docx")
    doc.save(output_path)
    return output_path


def _read_docx_content(docx_path: str) -> str:
    try:
        from docx import Document
    except Exception:
        return ""
    try:
        doc = Document(docx_path)
        full_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_content.append(paragraph.text.strip())
        for table_idx, table in enumerate(doc.tables):
            full_content.append(f"\n【表格{table_idx + 1}】")
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                full_content.append(" | ".join(row_data))
        return "\n\n".join(full_content)
    except Exception:
        return ""


def _generate_report_evaluation(report_text: str, config: RunnableConfig) -> str:
    try:
        model = get_llm(config)
        sys_prompt = _read_prompt("../prompts/summarizer_eval.md")
        prompt = ChatPromptTemplate.from_messages([
            ("system", sys_prompt),
            ("human", "{report}")
        ])
        messages = prompt.format_messages(report=report_text or "")
        resp = invoke_llm(
            model,
            messages,
            config=config,
            node="Summarizer",
            purpose="report_evaluation",
            json_mode=True,
        )
        return str(getattr(resp, "content", "")).strip()
    except Exception:
        return "报告生成完成。评价暂不可用。"


def summarizer(state: State, config: RunnableConfig, **kwargs):
    sections = _content_reorganizer(state)
    docx_path = _docx_report_generator(state, sections)

    if not docx_path:
        # 若无法生成 docx，降级生成 Markdown 报告
        os.makedirs("cache/report", exist_ok=True)
        md_path = os.path.join("cache/report", "final_report.md")
        try:
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(f"# {find_title(state)}\n\n")
                for sec in sections:
                    f.write(f"## {sec.get('title', '章节')}\n\n")
                    f.write(f"{sec.get('text', '')}\n\n")
        except Exception:
            pass
        docx_path = md_path

    report_text = _read_docx_content(docx_path)
    eval_text = _generate_report_evaluation(report_text, config)

    msg = AIMessage(content=eval_text)
    final_result = {
        "summary": "报告与评价已生成。",
        "attachments": [docx_path],
        "path": docx_path,
        "timestamp": datetime.now().isoformat(timespec="seconds"),
    }

    return {"messages": [msg], "final_result": final_result}
