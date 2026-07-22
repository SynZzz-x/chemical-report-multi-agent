import os
import re
import logging
import hashlib
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import io
from PIL import Image as PILImage
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def set_style(run, font_size=12, bold=False, italic=False, color=None):
    """
    设置 Run 的字体样式：
    - 西文：Arial
    - 中文：微软雅黑 (Microsoft YaHei)，兜底宋体 (SimSun)
    """
    run.font.name = 'Arial'
    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    if font_size:
        run.font.size = Pt(font_size)
    
    run.font.bold = bold
    run.font.italic = italic
    
    if color:
        run.font.color.rgb = color

def render_math_to_image(latex_str, fontsize=14, dpi=300):
    """
    使用 Matplotlib 将 LaTeX 公式渲染为图片 (BytesIO)
    """
    fig = None
    try:
        # 创建 Figure
        # 尺寸设得非常小，依赖 bbox_inches='tight' 扩展
        fig = plt.figure(figsize=(0.01, 0.01))
        fig.text(0, 0, f"${latex_str}$", fontsize=fontsize)
        
        buf = io.BytesIO()
        fig.savefig(buf, format='png', bbox_inches='tight', pad_inches=0.1, dpi=dpi)
        buf.seek(0)
        plt.close(fig)
        return buf
    except Exception as e:
        if fig:
            plt.close(fig)
        logger.warning(f"Math render failed for '{latex_str}': {e}")
        return None

def parse_inline_styles(text):
    """
    解析行内样式：公式($...$), 粗体 (**), 斜体 (*), 粗斜体 (***)
    返回一个列表，每个元素是字典：
    - 普通文本/粗体/斜体: {'type': 'text', 'content': str, 'bold': bool, 'italic': bool}
    - 公式: {'type': 'math', 'content': latex_str}
    """
    tokens = []
    
    # 定义正则模式
    # Group 1: Math ($...$)
    # Group 2: BoldItalic (***...***)
    # Group 3: Bold (**...**)
    # Group 4: Italic (*...*)
    pattern = re.compile(r'((?<!\\)\$(?:.+?)(?<!\\)\$)|(\*\*\*.*?\*\*\*)|(\*\*.*?\*\*)|(\*.*?\*)')
    
    last_idx = 0
    for match in pattern.finditer(text):
        start, end = match.span()
        
        # 添加匹配前的普通文本
        if start > last_idx:
            tokens.append({'type': 'text', 'content': text[last_idx:start], 'bold': False, 'italic': False})
            
        matched_text = match.group()
        
        # 判断匹配类型
        if match.group(1): # Math
            content = matched_text[1:-1] # 去掉 $
            tokens.append({'type': 'math', 'content': content})
        elif match.group(2): # BoldItalic
            content = matched_text[3:-3]
            tokens.append({'type': 'text', 'content': content, 'bold': True, 'italic': True})
        elif match.group(3): # Bold
            content = matched_text[2:-2]
            tokens.append({'type': 'text', 'content': content, 'bold': True, 'italic': False})
        elif match.group(4): # Italic
            content = matched_text[1:-1]
            tokens.append({'type': 'text', 'content': content, 'bold': False, 'italic': True})
            
        last_idx = end
        
    # 添加剩余文本
    if last_idx < len(text):
        tokens.append({'type': 'text', 'content': text[last_idx:], 'bold': False, 'italic': False})
        
    return tokens

def add_markdown_paragraph(doc, text, style='Normal', font_size=12, alignment=None):
    """
    向文档添加一个段落，支持行内 Markdown 样式解析
    """
    # 如果是标题样式，可能 docx 内部已有定义，我们这里先添加段落，再手动设置格式以确保字体正确
    p = doc.add_paragraph(style=style)
    
    if alignment is not None:
        p.alignment = alignment
        
    # 解析行内样式
    tokens = parse_inline_styles(text)
    
    # 确定基础是否加粗（例如标题默认加粗）
    base_bold = False
    if style.startswith('Heading'):
        base_bold = True
        
    for token in tokens:
        if token['type'] == 'text':
            run = p.add_run(token['content'])
            # 叠加样式
            is_bold = base_bold or token['bold']
            is_italic = token['italic']
            set_style(run, font_size=font_size, bold=is_bold, italic=is_italic)
            
        elif token['type'] == 'math':
            # 渲染公式图片并插入
            content = token['content']
            content_hash = hashlib.md5(content.encode('utf-8')).hexdigest()
            img_dir = os.path.abspath("./tmp/math_imgs")
            img_path = os.path.join(img_dir, f"{content_hash}.png")
            
            if not os.path.exists(img_path):
                 os.makedirs(img_dir, exist_ok=True)
                 img_buf = render_math_to_image(content, fontsize=14, dpi=300)
                 if img_buf:
                     with open(img_path, 'wb') as f:
                         f.write(img_buf.getvalue())
            
            if os.path.exists(img_path):
                try:
                    # 重新计算尺寸：
                    # 渲染时使用了 fontsize=14, dpi=300
                    # 目标是匹配当前文本 font_size (默认12)
                    
                    with PILImage.open(img_path) as img:
                        w_px, h_px = img.size
                        
                        # 原始渲染的自然高度 (Points)
                        # h_inch = h_px / dpi
                        # h_pt = h_inch * 72
                        h_pt_natural = (h_px / 300.0) * 72.0
                        
                        # 缩放比例
                        scale = font_size / 14.0 # 假设渲染用14，目标用 font_size
                        
                        target_height_pt = h_pt_natural * scale
                        target_width_pt = (w_px / 300.0) * 72.0 * scale
                        
                    run = p.add_run()
                    run.add_picture(img_path, width=Pt(target_width_pt), height=Pt(target_height_pt))
                    
                    # 垂直对齐调整
                    # 默认基线对齐。如果公式较高（如分数），会显得“靠上”（实际上是基线对齐导致顶部很高）
                    # 或者如果公式较矮（如 x ），也会基线对齐。
                    # 用户反馈“靠上”，意味着我们需要把它往下移。
                    # 简单的居中策略：
                    # 将图片的垂直中心与文本的垂直中心对齐。
                    # 文本高度约为 font_size。
                    # 偏移量 = (font_size - target_height_pt) / 2
                    # 如果 target_height_pt > font_size (高公式)，偏移量为负，向下移。
                    # 如果 target_height_pt < font_size (矮公式)，偏移量为正，向上移。
                    # 但考虑到数学公式基线通常比文本基线低一点点（为了对齐运算符），
                    # 我们可以稍微偏下一点。
                    
                    # 试用居中策略
                    offset_pt = (font_size - target_height_pt) / 2.0
                    
                    # 转换为 half-points (1/144 inch)
                    # w:position val is in half-points
                    offset_hp = int(offset_pt * 2)
                    
                    # 设置 w:position
                    if offset_hp != 0:
                         run.font.element.set(qn('w:position'), str(offset_hp))
                    
                except Exception as e:
                    logger.error(f"Failed to insert inline math image: {e}")
                    p.add_run(f"${content}$")
            else:
                p.add_run(f"${content}$")
        
    return p

def md_to_docx(content: str, output_path: str) -> str:
    """
    将 Markdown 转换为 DOCX
    """
    abs_output_path = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(abs_output_path), exist_ok=True)
    
    doc = Document()
    
    lines = content.split('\n')
    
    # 状态机
    in_code_block = False
    code_buffer = []
    
    # 标题自动编号堆栈
    # 规则：# 不编号，## 开始编号 (Index 0)
    heading_stack = []
    
    in_table = False
    table_buffer = []
    table_caption_text = None # 待处理的表格标题
    
    # Description Tag State
    in_description = False
    description_buffer = []
    pending_table_caption = None # 从 <description> 解析出的表格标题
    
    # Math Block State
    in_math_block = False
    math_buffer = []
    
    # 辅助：记录上一个添加的元素类型，用于图注判断
    # types: 'paragraph', 'image', 'table', 'code'
    last_element_type = None
    
    for line in lines:
        stripped_line = line.strip()
        
        # 1. 处理代码块
        if stripped_line.startswith('```'):
            if in_code_block:
                # 结束代码块
                in_code_block = False
                code_text = '\n'.join(code_buffer)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Courier New' # 代码字体
                run.font.size = Pt(10)
                p.style = 'No Spacing' # 紧凑样式
                # 可以加个边框或背景色，但 python-docx 原生支持较弱，暂略
                code_buffer = []
                last_element_type = 'code'
            else:
                # 如果之前在表格模式，先刷新表格
                if in_table:
                    _flush_table(doc, table_buffer, pending_table_caption)
                    in_table = False
                    table_buffer = []
                    pending_table_caption = None # 消费掉
                    last_element_type = 'table'
                in_code_block = True
            continue
            
        if in_code_block:
            code_buffer.append(line)
            continue
            
        # 1.5 处理公式块 ($$ ... $$)
        if stripped_line.startswith('$$'):
            if stripped_line == '$$':
                if in_math_block:
                    # End block
                    in_math_block = False
                    tex = '\n'.join(math_buffer)
                    math_buffer = []
                    
                    # 渲染并插入
                    img_buf = render_math_to_image(tex, fontsize=14, dpi=300)
                    if img_buf:
                        content_hash = hashlib.md5(tex.encode('utf-8')).hexdigest()
                        img_dir = os.path.abspath("./tmp/math_imgs")
                        img_path = os.path.join(img_dir, f"block_{content_hash}.png")
                        
                        if not os.path.exists(img_path):
                             os.makedirs(img_dir, exist_ok=True)
                             with open(img_path, 'wb') as f:
                                 f.write(img_buf.getvalue())
                        
                        try:
                            p = doc.add_paragraph()
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            run = p.add_run()
                            
                            with PILImage.open(img_path) as img:
                                w, h = img.size
                                # Max width 6 inches (approx page width with margins)
                                # Image rendered at 300 dpi
                                # Width in inches = w / 300
                                width_in_inches = w / 300.0
                                if width_in_inches > 6.0:
                                    run.add_picture(img_path, width=Inches(6.0))
                                else:
                                    run.add_picture(img_path, width=Inches(width_in_inches))
                                    
                            last_element_type = 'image'
                        except Exception as e:
                            logger.error(f"Failed to insert block math image: {e}")
                            doc.add_paragraph(f"$$ {tex} $$")
                            last_element_type = 'paragraph'
                    else:
                        doc.add_paragraph(f"$$ {tex} $$")
                        last_element_type = 'paragraph'
                else:
                    # Start block
                    in_math_block = True
                    math_buffer = []
                continue
            elif stripped_line.endswith('$$') and len(stripped_line) > 2:
                # Single line block $$...$$
                tex = stripped_line[2:-2]
                
                img_buf = render_math_to_image(tex, fontsize=14, dpi=300)
                if img_buf:
                    content_hash = hashlib.md5(tex.encode('utf-8')).hexdigest()
                    img_dir = os.path.abspath("./tmp/math_imgs")
                    img_path = os.path.join(img_dir, f"block_{content_hash}.png")
                    
                    if not os.path.exists(img_path):
                         os.makedirs(img_dir, exist_ok=True)
                         with open(img_path, 'wb') as f:
                             f.write(img_buf.getvalue())
                    
                    try:
                        p = doc.add_paragraph()
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run()
                        
                        with PILImage.open(img_path) as img:
                            w, h = img.size
                            width_in_inches = w / 300.0
                            if width_in_inches > 6.0:
                                run.add_picture(img_path, width=Inches(6.0))
                            else:
                                run.add_picture(img_path, width=Inches(width_in_inches))
                                
                        last_element_type = 'image'
                    except Exception as e:
                        logger.error(f"Failed to insert block math image: {e}")
                        doc.add_paragraph(f"$$ {tex} $$")
                        last_element_type = 'paragraph'
                else:
                    doc.add_paragraph(f"$$ {tex} $$")
                    last_element_type = 'paragraph'
                continue
        
        if in_math_block:
            math_buffer.append(line)
            continue

        # 2. 处理 <description> 标签 (图注/表注)
        # 必须在处理 Image/Table 之前处理
        if '<description>' in stripped_line or in_description:
            full_desc = ""
            # 如果是起始行
            if '<description>' in stripped_line:
                in_description = True
                start_idx = line.find('<description>') + len('<description>')
                if '</description>' in line:
                    end_idx = line.find('</description>')
                    content_desc = line[start_idx:end_idx].strip()
                    full_desc = content_desc
                    in_description = False
                    description_buffer = []
                else:
                    content_desc = line[start_idx:].strip()
                    if content_desc:
                        description_buffer.append(content_desc)
                    continue
            elif in_description:
                if '</description>' in stripped_line:
                    end_idx = line.find('</description>')
                    content_desc = line[:end_idx].strip()
                    if content_desc:
                        description_buffer.append(content_desc)
                    full_desc = " ".join(description_buffer).strip()
                    in_description = False
                    description_buffer = []
                else:
                    description_buffer.append(stripped_line)
                    continue
            
            # 处理提取到的 full_desc
            if not in_description:
                # 逻辑：
                # 1. 检查上一元素是否为图片
                if last_element_type == 'image':
                    # 添加图注（居中，灰色，小字）
                    p = doc.add_paragraph(full_desc)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in p.runs:
                        set_style(run, font_size=9, color=RGBColor(105, 105, 105), italic=False)
                    # 清除可能存在的 pending (以防万一)
                    pending_table_caption = None
                    last_element_type = 'caption'
                else:
                    # 解析为表格标题 (Pending)，等待下一个表格消费
                    pending_table_caption = full_desc
            
            # 如果当前行只包含标签，跳过后续处理
            if '<description>' in stripped_line and '</description>' in stripped_line:
                continue
            if '</description>' in stripped_line:
                continue

        # 2.5 处理分隔符 (---)
        if re.match(r'^[-*_]{3,}$', stripped_line):
            # 添加带有下边框的空段落模拟分割线
            p = doc.add_paragraph()
            p_element = p._p
            pPr = p_element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'auto')
            pBdr.append(bottom)
            pPr.append(pBdr)
            last_element_type = 'separator'
            continue

        # 3. 处理表格
        if stripped_line.startswith('|') and stripped_line.endswith('|'):
            if not in_table:
                # 只要进入表格模式，就暂存数据，不立即创建表格
                in_table = True
                # 如果有 pending 的标题，会在 _flush_table 时使用
            
            parts = [p.strip() for p in line.split('|')]
            if len(parts) > 2:
                row_data = parts[1:-1]
                # 检查分隔行 | --- | --- |
                if not all(set(c) <= set('-: ') for c in ''.join(row_data)):
                    table_buffer.append(row_data)
            continue
        else:
            if in_table:
                _flush_table(doc, table_buffer, pending_table_caption)
                in_table = False
                table_buffer = []
                pending_table_caption = None
                last_element_type = 'table'
        
        if not stripped_line:
            continue
            
        # 4. 处理图片 (![alt](src))
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', stripped_line)
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2)
            
            if os.path.exists(img_path):
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(5.5)) # 限制宽度
                    last_element_type = 'image'
                    
                    # 如果有 alt_text，可以作为备选图注，但如果有 <description> 标签，后者会追加在后面
                    # 为了避免重复，这里不自动将 alt_text 转为显式图注，除非用户没有用 <description>
                    # 但 md_to_pdf 中是直接添加了。为了保持一致性：
                    # md_to_pdf 逻辑：story.append(Paragraph(f"<i>图: {alt_text}</i>", caption_style))
                    # 这里我们可以先不加，等待下一行看有没有 <description>。
                    # 但如果下一行没有 <description> 怎么办？
                    # 这是一个 lookahead 问题。
                    # 简化策略：不使用 alt_text 作为显式图注，强制要求用户使用 <description> 或者自己在 markdown 里写。
                    # 或者：我们可以暂时不处理 alt_text，因为用户明确说了“约定输入markdown图和表的嵌入方式...则将标签对中间的文字解析为图注”。
                    # 这暗示只有标签对才是图注。alt text 可能只是为了 markdown 兼容。
                    pass 
                except Exception as e:
                    logger.warning(f"Failed to add image {img_path}: {e}")
                    doc.add_paragraph(f"[Image load failed: {img_path}]")
                    last_element_type = 'paragraph'
            else:
                doc.add_paragraph(f"[Image not found: {img_path}]")
                last_element_type = 'paragraph'
            continue
            
        # 5. 标题
        if stripped_line.startswith('# '):
            # Heading 1: 文档标题，不编号，重置堆栈
            heading_stack = []
            add_markdown_paragraph(doc, stripped_line[2:], style='Heading 1', font_size=16, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
            last_element_type = 'paragraph'
        elif stripped_line.startswith('## '):
            # Level 2 -> Index 0
            while len(heading_stack) <= 0:
                heading_stack.append(0)
            heading_stack[0] += 1
            heading_stack = heading_stack[:1]
            
            num_str = f"{heading_stack[0]}."
            add_markdown_paragraph(doc, f"{num_str} {stripped_line[3:]}", style='Heading 2', font_size=14)
            last_element_type = 'paragraph'
        elif stripped_line.startswith('### '):
            # Level 3 -> Index 1
            while len(heading_stack) <= 1:
                heading_stack.append(0)
            heading_stack[1] += 1
            heading_stack = heading_stack[:2]
            
            num_str = f"{heading_stack[0]}.{heading_stack[1]}"
            add_markdown_paragraph(doc, f"{num_str} {stripped_line[4:]}", style='Heading 3', font_size=12)
            last_element_type = 'paragraph'
        elif stripped_line.startswith('#### '):
            # Level 4 -> Index 2
            while len(heading_stack) <= 2:
                heading_stack.append(0)
            heading_stack[2] += 1
            heading_stack = heading_stack[:3]
            
            num_str = f"{heading_stack[0]}.{heading_stack[1]}.{heading_stack[2]}"
            add_markdown_paragraph(doc, f"{num_str} {stripped_line[5:]}", style='Heading 4', font_size=11)
            last_element_type = 'paragraph'
            
        # 6. 无序列表
        elif stripped_line.startswith('- ') or stripped_line.startswith('* '):
            add_markdown_paragraph(doc, stripped_line[2:], style='List Bullet')
            last_element_type = 'paragraph'
            
        # 7. 有序列表
        elif re.match(r'^\d+\.\s', stripped_line):
            match = re.match(r'^(\d+)\.\s+(.*)', stripped_line)
            content_text = match.group(2)
            add_markdown_paragraph(doc, content_text, style='List Number')
            last_element_type = 'paragraph'
            
        # 8. 普通文本
        else:
            add_markdown_paragraph(doc, stripped_line)
            last_element_type = 'paragraph'
            
    # 循环结束后，检查是否还在 table 中
    if in_table and table_buffer:
        _flush_table(doc, table_buffer, pending_table_caption)
        
    try:
        doc.save(abs_output_path)
        logger.info(f"DOCX generated successfully at {abs_output_path}")
    except Exception as e:
        logger.error(f"Failed to save DOCX: {e}")
        raise
        
    return abs_output_path

def _flush_table(doc, table_buffer, caption=None):
    if not table_buffer:
        return
        
    # 添加表注
    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in p.runs:
            set_style(run, font_size=10, bold=True)
            
    rows = len(table_buffer)
    cols = len(table_buffer[0]) if rows > 0 else 0
    
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    for r_idx, row_data in enumerate(table_buffer):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < cols:
                cell = row.cells[c_idx]
                # 清空默认段落（如果有）
                cell.text = ""
                # 使用 add_markdown_paragraph 往单元格添加内容，支持粗体等
                # 注意：cell.add_paragraph 会追加，cell.paragraphs[0] 是默认的
                p = cell.paragraphs[0]
                # 重新实现一个简单的 add_run logic for cell
                
                tokens = parse_inline_styles(cell_text)
                for token in tokens:
                    if token['type'] == 'text':
                        run = p.add_run(token['content'])
                        is_header = (r_idx == 0)
                        is_bold = token['bold'] or is_header # 表头默认加粗
                        set_style(run, font_size=10, bold=is_bold, italic=token['italic'])
                    elif token['type'] == 'math':
                        # Table cell math - try to insert image
                        content = token['content']
                        # Similar logic to add_markdown_paragraph math handling
                        # For simplicity, just insert text for now or try image?
                        # Let's try image if possible, but cell.paragraphs works same way
                        
                        # Just reuse add_markdown_paragraph logic?
                        # No, we already have p = cell.paragraphs[0]
                        # We need to manually add image run
                        
                        img_buf = render_math_to_image(content, fontsize=12, dpi=300)
                        if img_buf:
                             # Save temp
                             content_hash = hashlib.md5(content.encode('utf-8')).hexdigest()
                             img_dir = os.path.abspath("./tmp/math_imgs")
                             img_path = os.path.join(img_dir, f"cell_{content_hash}.png")
                             if not os.path.exists(img_path):
                                 os.makedirs(img_dir, exist_ok=True)
                                 with open(img_path, 'wb') as f:
                                     f.write(img_buf.getvalue())
                             
                             try:
                                  run = p.add_run()
                                  
                                  # 尺寸计算
                                  with PILImage.open(img_path) as img:
                                     w_px, h_px = img.size
                                     # 假设表内文字大小 10pt (参见 set_style 调用)
                                     cell_font_size = 10
                                     
                                     h_pt_natural = (h_px / 300.0) * 72.0
                                     scale = cell_font_size / 14.0 
                                     target_height_pt = h_pt_natural * scale
                                     target_width_pt = (w_px / 300.0) * 72.0 * scale
                                     
                                  run.add_picture(img_path, width=Pt(target_width_pt), height=Pt(target_height_pt))
                                  
                                  # 垂直对齐
                                  offset_pt = (cell_font_size - target_height_pt) / 2.0
                                  offset_hp = int(offset_pt * 2)
                                  if offset_hp != 0:
                                       run.font.element.set(qn('w:position'), str(offset_hp))
                                       
                             except Exception as e:
                                  logger.warning(f"Cell math error: {e}")
                                  p.add_run(f"${content}$")
                        else:
                             p.add_run(f"${content}$")
                    
                # 居中
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

if __name__ == "__main__":
    # Test
    md = r"""# Markdown 转 DOCX 测试报告 (含公式)
    
## 1. 字体测试
普通文本：这是 MiSans Normal 字体（在 DOCX 中应为微软雅黑）。
**粗体文本：这是 Bold 字体。**
*斜体文本：这是 Italic 字体。*
***粗斜体文本：Bold + Italic。***

## 2. 公式测试
### 2.1 行内公式
这是一个行内公式 $E=mc^2$ 测试。
这是一个稍微复杂的行内公式：$x = \frac{-b \pm \sqrt{b^2 - 4ac}}{2a}$。
看看对齐情况：$a^2+b^2=c^2$。

### 2.2 块级公式
下面是块级公式（单行）：
$$F = m \cdot a$$

下面是块级公式（多行）：
$$
\int_{a}^{b} f(x) dx = F(b) - F(a)
$$

## 3. 列表测试
### 无序列表
- 项目 A
- 项目 B

### 有序列表
1. 第一步：打开冰箱
2. 第二步：把大象装进去
3. 第三步：关上冰箱

## 4. 表格测试 (含公式)
<description>表1: 员工信息表 (含公式)</description>
| 姓名 | 指标 | 公式 | 备注 |
| --- | --- | --- | --- |
| 牛顿 | 力 | $F=ma$ | **经典** |
| 爱因斯坦 | 能量 | $E=mc^2$ | *质能* |
| 欧拉 | 恒等式 | $e^{i\pi} + 1 = 0$ | 美 |

## 5. 图片测试
![AltIgnored](test_image.png)
<description>图1: 测试图片-居中图注</description>

## 6. 代码测试
```python
def hello_world():
    print("Hello, Agent!")
    return True
```

## 7. 分隔符测试
---
(上面应该有一条线)
"""
    path = "./test_docx_full_v2.docx"
    print(f"Generating DOCX at {path}...")
    try:
        # 确保图片存在
        if not os.path.exists('test_image.png'):
            # 创建个假的
            from PIL import Image
            img = Image.new('RGB', (100, 50), color = (73, 109, 137))
            img.save('test_image.png')
            
        md_to_docx(md, path)
        print(f"Success! Check {path}")
    except Exception as e:
        print(f"Failed: {e}")
        import traceback
        traceback.print_exc()
